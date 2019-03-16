import yaml, pdb
from docx import Document
from string import Template

with open("config.yml", "r") as f:
    config_data = yaml.load(f)

doc = Document("templete.docx")
device_model = config_data["Device"]["Name"]
device_certification = config_data["Device"]["Description"]
date = {
    "reviewer": ("01", "07", "2018"),
    "approved": ("11", "07", "2018"),
    "release": ("13", "07", "2018")
}
data = {
    "device_model": device_model,
    "device_certification": device_certification,
    "rv_date": date["reviewer"][0],
    "rv_mon": date["reviewer"][1],
    "rv_year": date["reviewer"][2],
    "ap_date": date["approved"][0],
    "ap_mon": date["approved"][1],
    "ap_year": date["approved"][2],
    "re_date": date["release"][0],
    "re_mon": date["release"][1],
    "re_year": date["release"][2],
}


for i, paragraph in enumerate(doc.paragraphs):
    text = Template(paragraph.text).substitute(**data)
    print("style {}, text: {}".format(paragraph.style, text))
    paragraph.text = text

p = doc.add_paragraph("The table below details the build status of the DUT (Device under test) ")
p.add_run(device_model).bold = True

doc.save("test.docx")
