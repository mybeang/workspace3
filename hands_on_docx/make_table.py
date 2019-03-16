import pdb
from docx import Document
from docx.enum.style import WD_BUILTIN_STYLE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

string_data = """Index  Act_rx_cnt  Exp_rx_cnt Flow_name  Result
0            0           0     flow0    True
1           10           0     flow1   False
2           10           0     flow2   False
3           10           0     flow3   False
4           10           0     flow4   False
5           10           0     flow5   False
6           10           0     flow6   False
7           10           0     flow7   False
8            0           0     flow8    True
9            0           0     flow9    True
10           0           0    flow10    True
11           0           0    flow11    True
12           0           0    flow12    True
13           0           0    flow13    True
14          10          10    flow14    True"""

doc = Document()

header = ["Result", "Act_rx_cnt", "Exp_rx_cnt", "flow_name"]
data_list = string_data.splitlines()[1:]

records = list()

for data in data_list:
    items = data.split()[1:]
    records.append((items[3], items[0], items[1], items[2]))

table = doc.add_table(rows=1, cols=len(header))
table.style = 'Colorful List' #https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
for hdr_cell, hdr in zip(hdr_cells, header):
    hdr_cell.text = hdr
    hdr_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    par = hdr_cell.paragraphs[0]
    par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font = par.runs[0].font
    font.bold = True

for record in records:
    row_cells = table.add_row().cells
    for row_cell, text in zip(row_cells, record):
        row_cell.text = str(text)
        par = row_cell.paragraphs[0]
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        font = par.runs[0].font
        if text.lower() == 'true':
            font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
        elif text.lower() == 'false':
            font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

doc.add_page_break()
doc.save('table.docx')

