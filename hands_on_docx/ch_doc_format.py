from docx import Document
from docx.shared import Cm, Pt, RGBColor


DEFAULT_FONT = 'Consolas'
LOG_FONT = 8


class DocxLogs(object):
    def __init__(self):
        self.doc = Document()
        sections = self.doc.sections
        for section in sections:
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
        self.__report_string = ""

    def post_processing(self, log_string):
        splited_string = log_string.split(" ", 3)
        try:
            types = splited_string[2]
        except:
            types = ""
        return log_string, types

    def apply_font(self, string):
        _types_color = {
            'pass': "blue",
            'fail': "red",
            'debug': 'gray'
        }
        _string, _lv = self.post_processing(string)
        p = self.doc.add_paragraph()
        pfmt = p.paragraph_format
        pfmt.space_before = 1
        pfmt.space_after = 1
        pfmt.line_spacing = 1
        padd = p.add_run(_string)
        pfont = padd.font
        pfont.name = DEFAULT_FONT
        pfont.size = Pt(LOG_FONT)
        if _lv.lower() == 'report':
            pfont.bold = True
        elif _lv.lower() == ('setup' or 'test' or 'teardn'):
            pfont.italic = True
        elif _lv.lower() == 'fail':
            pfont.color.rgb = RGBColor(0xff, 0x00, 0x00)
        elif _lv.lower() == 'pass':
            pfont.color.rgb = RGBColor(0x00, 0x00, 0xff)
        elif _lv.lower() == 'debug':
            pfont.color.rgb = RGBColor(0x60, 0x60, 0x60)

    def save_file(self, filename):
        self.doc.add_page_break()
        self.doc.save(filename)


if __name__=="__main__":
    with open("test_log.log", 'r') as f:
        string_list = f.readlines()
    docx_logs = DocxLogs()
    for string in string_list:
        docx_logs.apply_font(string.strip())
    docx_logs.save_file("test_log.docx")

